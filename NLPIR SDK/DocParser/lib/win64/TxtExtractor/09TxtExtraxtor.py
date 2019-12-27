# encoding=utf-8
# version: python2.7
import codecs
import copy
import json
import os
import re
import sys
import time
import glob
import xlrd
import ctypes
import psutil
import chardet
import logging
import subprocess
import pptx
import docx
from lxml import etree
from queue import Queue
from xml.dom import minidom
from decimal import Decimal
from threading import Thread
from pdfminer.layout import *
from psutil import NoSuchProcess
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter


# 日志函数,指定日志文件名称,日志等级
def set_log(log_level=logging.INFO):
    logging.basicConfig(
        level=log_level,  # 定义输出到文件的log级别，
        format='%(asctime)s : %(levelname)s %(message)s',  # 定义输出log的格式
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    # formatter = logging.Formatter('%(asctime)s: %(levelname)s %(message)s')
    # files = logging.FileHandler(log_file)
    # files.setFormatter(formatter)
    # logging.getLogger("").addHandler(files)


# 重新定义write_xml方法
def fixed_writexml(self, writer, indent="", add_indent="", new_l=""):
    # indent = current indentation
    # add_indent = indentation to add to higher levels
    # newl = newline string
    writer.write(indent + "<" + self.tagName)

    attrs = self._get_attributes()
    a_names = attrs.keys()
    a_names.sort()

    for a_name in a_names:
        writer.write(" %s=\"" % a_name)
        minidom._write_data(writer, attrs[a_name].value)
        writer.write("\"")
    if self.childNodes:
        if len(self.childNodes) == 1 \
                and self.childNodes[0].nodeType == minidom.Node.TEXT_NODE:
            writer.write(">")
            self.childNodes[0].writexml(writer, "", "", "")
            writer.write("</%s>%s" % (self.tagName, new_l))
            return
        writer.write(">%s" % new_l)
        for node in self.childNodes:
            if node.nodeType is not minidom.Node.TEXT_NODE:
                node.writexml(writer, indent + add_indent, add_indent, new_l)
        writer.write("%s</%s>%s" % (indent, self.tagName, new_l))
    else:
        writer.write("/>%s" % new_l)


# 清理程序运行后残留的word，ppt等后台进程，防止下次运行出现错误
def clear_process():
    try:
        pids = psutil.pids()
        for pid in pids:
            # 获取所有进程
            p = psutil.Process(pid)
            # 寻找word进程进行清理
            if "WORD" in p.name():
                # print("pid-%d,     pname-%s" %(pid,  p.name()))
                # 杀死进程
                psutil.Process(pid).terminate()
            # 寻找ppt进程进行清理
            if "POWERPNT" in p.name():
                # print("pid-%d,     pname-%s" % (pid, p.name()))
                # 杀死进程
                psutil.Process(pid).terminate()
    except NoSuchProcess:
        pass


# 获取文件大小
def get_size(route):
    s = os.path.getsize(route)
    if s / 1024 < 1024:
        # print("{}K".format(s / 1024))
        if s < 1024:
            return "1K"
        if s % 1024 > 0:
            return "{}K".format(s / 1024 + 1)
        else:
            return "{}K".format(s / 1024)
    else:
        d = Decimal("{}".format(float(s / 1024) / 1024)).quantize(Decimal('0.00'))
        # print("{}M".format(d))
        return "{}M".format(d)


# 获取每个文件转换信息保存成字典并存进队列中
def save_log_dict(dict_info_list):
    # 保存相关转换信息, 参数列表中 1.文件名 2.原路径 3.保存路径 4.源文件大小 5.保存文件大小 6.文件格式
    extract_dict = dict()
    extract_dict[LogConstant.TIME] = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(time.time()))
    extract_dict[LogConstant.NAME] = dict_info_list[0]
    extract_dict[LogConstant.FROM_ROUTE] = dict_info_list[1]
    extract_dict[LogConstant.TO_ROUTE] = dict_info_list[2]
    extract_dict[LogConstant.FROM_SIZE] = dict_info_list[3]
    extract_dict[LogConstant.TO_SIZE] = dict_info_list[4]
    extract_dict[LogConstant.FORMAT] = dict_info_list[5]
    log_queue.put(extract_dict)  # 将转换信息存到log_queue队列中，用于单独提取
    # logging.info(extract_dict)


# 手动程序关闭
def controller():
    logging.info("按回车键可强制结束!") # .decode("utf-8").encode("gb2312"))
    try:
        raw_input()
    except Exception:
        pass
    # 清理残余进程
    clear_process()
    on_exit()


# 退出函数
def on_exit():
    os._exit(0)


set_log(log_level=logging.INFO)
# 重写write_xml方法
minidom.Element.writexml = fixed_writexml
log_queue = Queue()  # 定义文件转换信息的队列


class LogConstant(object):
    """
        日志文件保存相关信息。
    """
    # 日志路径
    LOG_DIR = "logs"  # 日志保存文件夹
    LOG_FILE_NAME = time.strftime("%Y-%m-%d", time.localtime(time.time()))  # 获取当天时间,当作日志文件名

    # xml标签
    TOTAL_CONTENT = "total_content"  # 最外层标签，根节点，保存所有内容
    FILE_INFO = "file_info"    # 一级节点，每个文件所有转换信息
    NAME = "name"              # 二级节点，转换的文件名
    FROM_ROUTE = "from_route"  # 二级节点，原始文件路径
    FROM_SIZE = "from_size"    # 二级节点，原始文件大小
    TO_ROUTE = "to_route"      # 二级节点，生成文件路径
    TO_SIZE = "to_size"        # 二级节点，生成文件大小
    TIME = "extract_time"              # 二级节点，保存当前时间
    FORMAT = "format"          # 二级节点，保存文件格式


class SaveXmlLog(object):
    def __init__(self):
        self.none = 0

    @staticmethod
    def create_node(dom, item, content):
        # --创建一个二级节点，以及内容添加到一级节点下
        second_node = dom.createElement(item[0])
        try:
            node_con = dom.createTextNode(item[1].decode("utf-8"))
        except Exception:
            node_con = dom.createTextNode(item[1].decode("gbk"))
        second_node.appendChild(node_con)
        # 将二级标签添加到一级标签下
        content.appendChild(second_node)

    def save_xml_log(self):
        # 将内容保存成xml文件
        while True:
            item = log_queue.get()
            log_queue.task_done()
            if item is "None":
                break
            else:
                # 在已有的xml中添加内容
                dom = minidom.parse("{}/{}.xml".format(LogConstant.LOG_DIR, LogConstant.LOG_FILE_NAME))
                root = dom.documentElement
                # --创建一个一级节点，并添加到root下
                content = dom.createElement(LogConstant.FILE_INFO)
                root.appendChild(content)
                for i in item.items():
                    # todo: 循环遍历字典，可能导致保存的内容顺序是不固定的，后续看需求进行修改
                    # 创建一个二级节点，以及内容添加到一级节点下
                    self.create_node(dom, i, content)
                # -- 写进文件,如果出现了unicode，指定文件的编码
                f = codecs.open("{}/{}.xml".format(LogConstant.LOG_DIR, LogConstant.LOG_FILE_NAME), 'w', 'utf-8')
                dom.writexml(f, addindent='  ', newl='\n', encoding='utf-8')
                f.close()

    @staticmethod
    def create_xml():
        # 1.判断是否有logs文件夹，若没有则生成log日志文件夹
        if not os.path.exists(LogConstant.LOG_DIR):
            os.mkdir(LogConstant.LOG_DIR)
        if not os.path.exists("{}/{}.xml".format(LogConstant.LOG_DIR, LogConstant.LOG_FILE_NAME)):
            # 创建xml初始文件
            doc = minidom.Document()
            # 创建最外层标签
            total = doc.createElement(LogConstant.TOTAL_CONTENT)
            doc.appendChild(total)
            # 创建文件写入根目录内容
            f = open(r"{}/{}.xml".format(LogConstant.LOG_DIR, LogConstant.LOG_FILE_NAME), "w")
            f.write(doc.toprettyxml(indent="  "))
            f.close()

    def run(self):
        # 阻塞模式，保证不会因空队列而退出
        log_queue.join()
        # 1. 新建文件夹和xml文件并且存储内容，使用队列
        self.create_xml()
        # 2. 将要保存的日志信息保存进xml文件中
        self.save_xml_log()


class PdfExtractor:
    """pdf文本提取类"""

    def __init__(self, route, save_route):
        # 获取文件路径及名称
        self.route = route  # .decode("utf-8").encode("gbk")
        # 获取保存文件路径
        self.save_route = save_route

    def all_pdf(self):
        # 获取当前路径下所有pdf文件路径
        for d in glob.glob(self.route + r'/*.pdf'):
            logging.info(d)

    def change_type(self, file_name):
        # 判断传递参数路径是否是个文件
        if self.save_route[-3:] == "txt":
            route = self.save_route  # 是带文件路径
            save_only_route = "/".join(self.save_route.split("/")[:-1])
        else:
            route = self.save_route + r'/{}.txt'.format(file_name)  # 不是文件，拼接成路径文件
            save_only_route = self.save_route
        return route, save_only_route

    def exchange(self):
        # 提取内容并保存进txt中
        try:
            fp = open(self.route, 'rb')
            # 获取文件名
            file_name = ".".join(self.route.split("/")[-1].split(".")[:-1])
            file_all_name = self.route.split("/")[-1]
            # 获取文件纯路径
            only_route = "/".join(self.route.split("/")[:-1])
            logging.info("正在提取" + " {}".format(file_all_name))
            # 来创建一个pdf文档分析器
            parser = PDFParser(fp)
            # 创建一个PDF文档对象存储文档结构
            document = PDFDocument(parser)
            # 检查文件是否允许文本提取
            if document.is_extractable:
                rsrcmgr = PDFResourceManager()  # 创建一个PDF资源管理器对象来存储共赏资源
                laparams = LAParams()  # 设定参数进行分析
                device = PDFPageAggregator(rsrcmgr, laparams=laparams)  # 创建一个PDF设备对象
                interpreter = PDFPageInterpreter(rsrcmgr, device)  # 创建一个PDF解释器对象
                route, save_only_route = self.change_type(file_name)
                # 处理每一页
                for page in PDFPage.create_pages(document):
                    interpreter.process_page(page)
                    layout = device.get_result()  # 接受该页面的LTPage对象
                    for x in layout:
                        if isinstance(x, LTTextBoxHorizontal):
                            with open(route, 'a') as f:
                                f.write(x.get_text().encode('utf-8') + '\n')
                        # end for
                    # end for
                # 按顺序将信息存入列表中 1.文件名 2.原路径 3.保存路径 4.源文件大小 5.保存文件大小 6.文件格式
                dict_info_list = [file_name, only_route, save_only_route, get_size(self.route),
                                  get_size(route), "pdf"]
                # 保存信息
                save_log_dict(dict_info_list)
                # 显示当前文件内容提取结束
                logging.info("{} ".format(file_all_name) + "提取结束!")
        except Exception as e:
            logging.error("{}{}".format("文件抽取失败！错误：", e))

    def run(self):
        # 运行转换文件
        self.exchange()


class PptExtractor:
    """ppt提取类"""

    def __init__(self, route, save_route):
        # 获取文件路径
        self.route = route  # .decode("utf-8").encode("gbk")
        # 获取保存文件路径
        self.save_route = save_route

    def all_ppt(self):
        # 获取当前路径下所有pdf文件路径
        for d in glob.glob(self.route + r'/*.pptx'):
            # self.route_queue.put(d)
            logging.info(d)

    def change_type(self, file_name):
        # 判断传递参数路径是否是个文件
        if self.save_route[-3:] == "txt":
            route = self.save_route  # 是带文件路径
            save_only_route = "/".join(self.save_route.split("/")[:-1])
        else:
            route = self.save_route + r'/{}.txt'.format(file_name)  # 不是文件，拼接成路径文件
            save_only_route = self.save_route
        return route, save_only_route

    def exchange(self):
        # 获取文件路径及名称
        try:
            # 获取文件名
            file_name = ".".join(self.route.split("/")[-1].split(".")[:-1])
            file_all_name = self.route.split("/")[-1]
            only_route = "/".join(self.route.split("/")[:-1])
            logging.info("正在提取" + " {}".format(file_all_name))
            route, save_only_route = self.change_type(file_name)
            # 打开文件
            ppt_file = pptx.Presentation(self.route)
            total_con = ""
            for shape in ppt_file.slides:
                data_s = shape.element.xml
                con_list = re.findall(r"<a:t>(.*)</a:t>", data_s)
                partial_content = ""
                for i in con_list:
                    partial_content += i
                total_con += partial_content + "\t\n\n"
            f = file(route, "w")
            f.write(total_con.encode('utf-8'))
            f.close()
            # 显示当前文件内容提取结束
            logging.info("{} ".format(file_all_name) + "提取结束!")
            # 按顺序将信息存入列表中 1.文件名 2.原路径 3.保存路径 4.源文件大小 5.保存文件大小 6.文件格式
            dict_info_list = [file_name, only_route, save_only_route, get_size(self.route),
                              get_size(route), "pptx"]
            # 保存信息
            save_log_dict(dict_info_list)
        except Exception:
            logging.info("{}-{}".format("抽取失败！", self.route))

    def run(self):
        # 运行
        self.exchange()


class WordExtractor:
    """word文档内容提取类"""

    def __init__(self, route, save_route):
        # 获取文件路径
        self.route = route  # .decode("utf-8").encode("gbk")
        # 获取保存文件路径
        self.save_route = save_route

    def all_word(self):
        # 获取当前路径下所有pdf文件路径
        for d in glob.glob(self.route + r'/*.doc'):
            # self.route_queue.put(d)
            logging.info(d)
        for d in glob.glob(self.route + r'/*.docx'):
            # self.route_queue.put(d)
            logging.info(d)

    def change_type(self, file_name):
        # 判断传递参数路径是否是个文件
        if self.save_route[-3:] == "txt":
            route = self.save_route  # 是带文件路径
            save_only_route = "/".join(self.save_route.split("/")[:-1])
        else:
            route = self.save_route + '/{}.txt'.format(file_name)  # 不是文件，拼接成路径文件
            save_only_route = self.save_route
        return route, save_only_route

    def exchange(self):
        try:
            # 获取文件名
            file_name = ".".join(self.route.split("/")[-1].split(".")[:-1])
            file_all_name = self.route.split("/")[-1]
            file_type = self.route.split("/")[-1].split(".")[-1]
            # 获取文件纯路径
            only_route = "/".join(self.route.split("/")[:-1])
            logging.info("正在提取" + " {}".format(file_all_name))
            # 提取文件内容
            name, ext = os.path.splitext(self.route)

            if ext != '.doc' and ext != '.docx':
                return None
            route, save_only_route = self.change_type(file_name)
            if ext == ".docx":
                dop = docx.Document(self.route)
                dop_text = '\n'.join([paragraph.text for paragraph in dop.paragraphs])
                with open(route, "w") as f:
                    f.write(dop_text.encode("utf-8"))
            else:
                dop = subprocess.check_output(["antiword", self.route]).decode("utf-8")
                with open(route, "w") as f:
                    f.write(dop.encode("utf-8"))
                # 按顺序将信息存入列表中 1.文件名 2.原路径 3.保存路径 4.源文件大小 5.保存文件大小 6.文件格式
                dict_info_list = [file_name, only_route, save_only_route, get_size(self.route),
                                  get_size(route), file_type]
                # 保存信息
                save_log_dict(dict_info_list)
                # 显示当前文件内容提取结束
                logging.info("{} ".format(file_all_name) + "提取结束!")
        except Exception:
                logging.error("{} {}".format("提取失败！", self.route))

    def run(self):
        # 运行
        self.exchange()


class ExcelExtractor:
    """excel内容提取类"""

    def __init__(self, route, save_route):
        # 获取文件路径
        self.route = route  # .decode("utf-8").encode("gbk")
        # 获取保存文件路径
        self.save_route = save_route

    def all_excel(self):
        # 获取当前路径下所有excel文件路径
        for d in glob.glob(self.route + r'/*.xlsx'):
            logging.info(d)
            # self.route_queue.put(d)

    def change_type(self, file_name):
        # 判断传递参数路径是否是个文件
        if self.save_route[-3:] == "txt":
            route = self.save_route  # 是带文件路径
            save_only_route = "/".join(self.save_route.split("/")[:-1])
        else:
            route = self.save_route + r'/{}.txt'.format(file_name)  # 不是文件，拼接成路径文件
            save_only_route = self.save_route
        return route, save_only_route

    def exchange(self):
        # 获取文件路径
        try:
            # 获取文件名
            file_name = ".".join(self.route.split("/")[-1].split(".")[:-1])
            file_all_name = self.route.split("/")[-1]
            file_type = self.route.split("/")[-1].split(".")[-1]
            # 获取文件纯路径
            only_route = "/".join(self.route.split("/")[:-1])
            logging.info("正在抽取" + " {}".format(file_all_name))
            route, save_only_route = self.change_type(file_name)
            data = xlrd.open_workbook(self.route)
            sqlfile = open(route, "a")  # 文件读写方式是追加
            table = data.sheets()[0]  # 表头
            n_rows = table.nrows  # 行数
            # n_cols = table.ncols  # 列数
            # col_names = table.row_values(0)  # 某一行数据
            for ro_num in range(n_rows):
                row = table.row_values(ro_num)
                values = ""
                for d in row:
                    values = values + json.dumps(d, ensure_ascii=False) + ","
                    values = values.replace('"', "")
                # 拼接行数据变成字符串，按逗号分隔
                values = values.encode("utf-8")
                sqlfile.writelines(values + "\r\n")  # 将字符串写入新文件
            # 关闭写入的文件
            sqlfile.close()
            # 按顺序将信息存入列表中 1.文件名 2.原路径 3.保存路径 4.源文件大小 5.保存文件大小 6.文件格式
            dict_info_list = [file_name, only_route, save_only_route, get_size(self.route),
                              get_size(route), file_type]
            # 保存信息
            save_log_dict(dict_info_list)
            # 显示当前文件内容提取结束
            logging.info("{} ".format(file_all_name) + "抽取结束!")
        except Exception:
            logging.error("{} {}".format("抽取失败！", self.route))

    def run(self):
        self.exchange()


class GeneralExtract:
    """
    通用提取器，可提取所有能用记事本打开且不乱码的文件
    params：要抽取内容的文件夹：extra_dir = r"C:\Users\Administrator\Desktop\desk\统计end".decode("utf-8").encode("gbk")
            要保存到的文件夹：  defined_dir = r"C:\Users\Administrator\Desktop\spider\NewFiles"

    """

    def __init__(self, extra_dir, defined_dir):
        self.spider_dir = extra_dir  # 提取文件最开始路径
        self.defined_dir = defined_dir  # 要保存的文件路径
        # 提取文件的最开始文件夹
        self.start_dir = self.spider_dir.split("/")[-1]
        self.queue = Queue()  # 队列，保存每个文件夹和相关信息
        self.file_queue = Queue()  # 队列， 保存每个文件和相关信息
        self.i = 0

    @staticmethod
    def code_exchange(file_content):
        # 将读取到的内容按照对应的编码格式进行解码，返回解码后的内容
        encoding = chardet.detect(file_content)["encoding"]
        if encoding is not None:
            # logging.info("文件读取到的内容的所有编码信息：{}".format(chardet.detect(file_content)))
            # logging.info("文件内容的编码格式是：{}".format(chardet.detect(file_content)))
            if "2312" in encoding:
                content = file_content.decode("gb18030")
            else:
                content = file_content.decode(encoding)
        else:
            content = file_content
        return content

    @staticmethod
    def get_single_file(name_p, routes, save_routes):
        if name_p == "pdf":
            pdf = PdfExtractor(routes, save_routes)
            pdf.run()
        elif name_p == "doc" or name_p == "docx":
            word = WordExtractor(routes, save_routes)
            word.run()
        elif name_p == "pptx":
            ppt = PptExtractor(routes, save_routes)
            ppt.run()
        elif name_p == "xlsx" or name_p == "xls":
            excel = ExcelExtractor(routes, save_routes)
            excel.run()
        else:
            logging.error("转换文件格式错误！")

    def clear_content(self, html_content):
        response = html_content
        # 获取内容解码异常处理
        try:
            try:
                html = etree.HTML(response.decode("utf-8"))
            except UnicodeDecodeError:
                html = etree.HTML(response.decode("gbk"))
        except UnicodeDecodeError:
            html = etree.HTML(response)
        # 对获取的内容进行编码
        try:
            try:
                response = response.decode("utf-8")
            except UnicodeDecodeError:
                response = response.decode("gbk")
        except UnicodeDecodeError:
            pass
        # 获取所有script  iframe   style标签中的内容
        js_str = html.xpath("//script/text()")
        if_str = html.xpath("//ifram/text()")
        style_str = html.xpath("//style/text()")
        # 将所有script  iframe  style标签中的内容替换成空内容
        response = self.js_replace(js_str, response)
        response = self.js_replace(if_str, response)
        response = self.js_replace(style_str, response)
        # 解析
        html = etree.HTML(response)
        return html

    @staticmethod
    def js_replace(ge_str, response):
        for str_s in ge_str:
            try:
                response = response.replace(str_s, "")
            except Exception:
                pass
        return response

    def extra_content(self, html_content, file_s):
        detail = None
        try:
            html = self.clear_content(html_content)
            # 提取页面所有内容
            detail = html.xpath("//html")[0].xpath("string(.)")
        except Exception:
            logging.error("{} {}".format("抽取失败".decode("utf-8").encode("gbk"), file_s))
        return detail

    def judge_type(self, root, get_dir, file_s, real_name):
        # 获取文件后缀名，判断文件类型
        type_f = file_s.split(".")[-1] if len(file_s.split(".")) > 1 else "none"
        if type_f in ["doc", "pdf", "docx", "xls", "xlsx", "pptx"]:
            # 开始路径
            routes = os.path.abspath(root) + "/{}".format(file_s)
            # 保存路径
            save_routes = self.defined_dir + "/" + self.start_dir + get_dir + "/" + "{}.txt".format(real_name)
            self.get_single_file(type_f, routes, save_routes)
        elif type_f in "html":
            # 读取文件内容
            with open(os.path.abspath(root) + "/{}".format(file_s), "rb") as f:
                f_content = f.read()
            real_content = self.extra_content(f_content, file_s)
            if real_content is not None:
                self.write_files(root, get_dir, file_s, real_name, real_content)
        elif type_f in ["png", "jpeg", "jpg", "gif", "none"]:
            logging.error("{} {}".format("无法读取的文件", file_s))
        else:
            # 读取文件内容
            with open(os.path.abspath(root) + "/{}".format(file_s), "r") as f:
                f_content = f.read()
            try:
                # 对获取到的内容进行相应的解码，返回解码后的内容
                real_content = self.code_exchange(f_content)
            except UnicodeDecodeError:
                real_content = f_content
            logging.info("{} {}".format('正在提取', file_s))
            self.write_files(root, get_dir, file_s, real_name, real_content)

    def write_files(self, root, get_dir, file_s, real_name, real_content):
        try:
            # 将文件内容写进txt文件
            with open(self.defined_dir + "/" + self.start_dir + get_dir + "/" + "{}.txt"
                      .format(real_name), "w") as w:
                # todo:许多没有汉字的文件内容，提取编码之后还是处于未编码状态，不知道为什么
                w.write(real_content.encode("utf-8"))
                logging.info("{} {} ".format(file_s, '提取结束！'))
            file_type = file_s.split(".")[-1] if len(file_s.split(".")) > 1 else ""
            root_file = os.path.abspath(root) + "/{}".format(file_s)
            save_only_route = self.defined_dir + "/" + self.start_dir + get_dir
            route = self.defined_dir + "/" + self.start_dir + get_dir + "/" + "{}.txt".format(real_name)
            # 按顺序将信息存入列表中 1.文件名 2.原路径 3.保存路径 4.源文件大小 5.保存文件大小 6.文件格式
            dict_info_list = [real_name, root, save_only_route, get_size(root_file), get_size(route), file_type]
            # 保存信息
            save_log_dict(dict_info_list)
        except UnicodeDecodeError:
            # 判断写入错误的文件是否保存了空文件，如果存在，则删除
            if os.path.exists(self.defined_dir + "/" + self.start_dir + get_dir + "/" + "{}.txt"
                              .format(real_name)):
                os.remove(self.defined_dir + "/" + self.start_dir + get_dir + "/" + "{}.txt".format(real_name))
            logging.error("{} {}".format('无法保存的文件',
                                         os.path.abspath(root) + "/" + file_s))

    def put_content_to_queue(self):
        dir_list = os.walk(self.spider_dir)
        for tuple_con in dir_list:
            # 将获取到的信息存进队列
            self.queue.put(copy.deepcopy(tuple_con))

    def get_content_from_queue(self):
        while True:
            if self.queue.empty():
                break
            else:
                tuple_con = self.queue.get()
                root = tuple_con[0]  # 当前目录路径(绝对路径)
                # dirs = tuple_con[1]  # 当前路径下所有子目录
                files = tuple_con[2]  # 当前路径下所有非目录子文件
                # 获取当前文件夹下面的子目录下相对路径
                get_dir = os.path.abspath(root).split(self.start_dir)[-1]
                # 在自定义文件夹中创建对应目录
                if not os.path.exists(self.defined_dir + "/" + self.start_dir + get_dir):
                    try:
                        os.makedirs(self.defined_dir + "/" + self.start_dir + get_dir)
                    except Exception:
                        logging.error("文件夹已经存在！")
                # 创建列表保存文件相关信息
                for file_s in files:
                    file_lists = [file_s, root, get_dir]
                    self.file_queue.put(copy.deepcopy(file_lists))
                self.queue.task_done()

    def get_save_content(self):
        while True:
            if self.file_queue.empty():
                self.i += 1
                # logging.info("****************************************{}".format(self.i))
                if self.i == 15:
                    logging.info("内容全部提取结束")
                    # 清理残余进程
                    clear_process()
                    while True:
                        if log_queue.empty():
                            time.sleep(1)
                            on_exit()
                        else:
                            continue
                break
            else:
                # 从队列中获取所有文件
                file_list = self.file_queue.get()
                file_s = file_list[0]
                root = file_list[1]
                get_dir = file_list[2]
                if file_s[:2] not in "~$":
                    # 提取文件名称，去除后缀名
                    real_name = ".".join(file_s.split(".")[:-1])
                    self.judge_type(root, get_dir, file_s, real_name)
                    self.file_queue.task_done()

    def run(self):
        thread_list = list()
        # 1.将文件夹中每个子目录所有信息保存进队列
        self.put_content_to_queue()
        # 2.开启多线程，获取每个子目录下的文件所有信息，并保存进第二个队列
        for i in range(5):
            get_thread = Thread(target=self.get_content_from_queue)
            thread_list.append(get_thread)
        for t in thread_list:
            t.start()
        thread_list_2 = list()
        # 3.判断保存子目录队列，当为空时，开启多线程保存文件
        while True:
            if self.queue.empty:
                for j in range(15):
                    file_thread = Thread(target=self.get_save_content)
                    thread_list_2.append(file_thread)
                for s in thread_list_2:
                    s.start()
                break
            else:
                continue


class Main:
    """
                                      1                2                              3
    参数格式：python 09TxtExtraxtor.py -f "C:\Users\Administrator\Desktop\ss.doc"  "C:\Users\Administrator\Desktop"
                            -f 表示转换的是一个文件    第一个路径是要转换的文件路径  第二个路径是要保存的位置，默认同前
                                    1                  2                              3
             python 09TxtExtraxtor.py -d "C:\Users\Administrator\Desktop"  "C:\Users\Administrator\Desktop"
                    -d 表示要转换文件夹里面的所有文件   要转换的文件夹路径    要保存的文件夹路径   默认当前文件夹
    """

    def __init__(self):
        # 通过加参数的形式获取要转换的文件路径
        self.routes = os.path.abspath(sys.argv[2])

    def extract_single_file(self):
        # 通过加参数的形式获取要转换的文件路径
        # routes = os.path.abspath(sys.argv[1])
        # 判断是否是一个文件，若是文件则正常转换，若不是则停止程序
        if not os.path.isfile(self.routes):
            logging.error("没有此文件，不能转换！")
            time.sleep(1)
            logging.info("程序结束！")
            time.sleep(1)
            on_exit()
        # os.path.abspath(path) 相对路径转变成绝对路径
        routes = os.path.abspath(self.routes)
        # 路径
        if routes[0] == '"' or routes[0] == "'":
            routes = routes[1:-1]
        # 异常判断，如果不输入第三个参数，默认路径和转换文件路径同级
        try:
            save_routes = os.path.abspath(sys.argv[3])
        except IndexError:
            save_routes = "/".join(routes.split("/")[:-1])
        # 文件名称
        name = routes.split("/")[-1]
        # 文件后缀名称
        name_p = name.split(".")[-1]
        if name_p == "pdf":
            pdf = PdfExtractor(routes, save_routes)
            pdf.run()
        elif name_p == "doc" or name_p == "docx":
            word = WordExtractor(routes, save_routes)
            word.run()
        elif name_p == "pptx":
            ppt = PptExtractor(routes, save_routes)
            ppt.run()
        elif name_p == "xlsx" or name_p == "xls":
            excel = ExcelExtractor(routes, save_routes)
            excel.run()
        else:
            logging.error("转换文件格式错误！")
        # 提取结束打印结束信息
        time.sleep(2)
        logging.info("程序结束")
        time.sleep(2)
        on_exit()

    def extract_dirs(self):
        extra_dir = self.routes  # 提取文件夹中的内容
        try:
            defined_dir = os.path.abspath(sys.argv[3])  # 获取第三个参数（文件保存路径）
        except IndexError:
            logging.error("缺少参数，没找到保存路径！")
            defined_dir = ""
            on_exit()
        g_extract = GeneralExtract(extra_dir, defined_dir)
        g_extract.run()

    def run(self):
        try:
            # 调用dll文件获取授权函数
            dll = ctypes.CDLL("./libLicense.so")
            # 调用dll中函数，返回1表示授权成功，否则授权失败
            result = dll.IsValidLicense("TxtExtractor.user", "TxtExtractor",0,0)
        except Exception:
            result =0
        # 判断是否授权成功
        if result == 1:
            logging.info("授权合法！")
            time.sleep(1)
            logging.info("开始提取！") 
            # 判断要抽取的是一个文件还是文件夹中内容
            type_f = sys.argv[1]
            if type_f == "-f":
                # 要抽取一个文件
                logging.info("抽取文件！") 
                self.extract_single_file()
            elif type_f == "-d":
                # 要抽取文件夹中的所有文件
                logging.info("抽取文件夹下所有文件！") 
                self.extract_dirs()
            else:
                logging.error("参数错误")
        else:
            logging.info("授权非法，无法抽取！")
            # 提取结束打印结束信息
            time.sleep(2)
            logging.info("程序结束")
            time.sleep(2)
            on_exit()


if __name__ == '__main__':
    # 创建保存xml日志线程
    xml_log = SaveXmlLog()
    controller_log = Thread(target=xml_log.run)
    controller_log.start()
    # 创建线程，手动控制程序强制关闭
    controller_time = Thread(target=controller)
    controller_time.start()
    # 创建主函数对象
    main = Main()
    controller_main = Thread(target=main.run)
    controller_main.start()
